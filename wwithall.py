import os
import docx
import docxpy
import re
from docx.shared import Cm

files=[]

for f in os.listdir('docs'):
    if(f.endswith('.docx')):
        files.append(f)
for f in files:
        doc=docx.Document('docs/'+f)
        for sec in doc.sections:
                sec.top_margin=Cm(3)
                sec.bottom_margin=Cm(1)
                sec.left_margin=Cm(2.5)
                sec.right_margin=Cm(1.5)
        dtables=0
        for t in  range(0,len(doc.tables)):
                doc.tables[t].allow_autofit=True
                dtables+=1
                doc.tables[dtables-1].width=Cm(10)
        for p in range(0,len(doc.paragraphs)):
                if(re.search(r'salgysy',doc.paragraphs[p].text)):
                        oldtext=doc.paragraphs[p].text
                        if 'Daşoguz welaýatynyň' not in oldtext:
                                ttext=oldtext.split(':')
                                newtext=ttext[0]+':'+'Daşoguz welaýatynyň'+ ttext[1].split('nyň')[0]+ttext[1].split('nyň')[1]
                                doc.paragraphs[p].text=newtext
                                doc.paragraphs[p].style.font.size=12700*16
                                doc.paragraphs[p].style.font.name='Times New Roman'
                                doc.paragraphs[p].style.font.bold=True
                if(re.search(r'Tel',doc.paragraphs[p].text)):
                        oldtext=doc.paragraphs[p].text
                        if '(ykjam)' not in oldtext:
                                ttext=oldtext.split('.')
                                if(len(ttext)>1):
                                        newtext='Telefony:'+ttext[1]+'(ykjam)'
                                        doc.paragraphs[p].text=newtext
                                        doc.paragraphs[p].style.font.size=12700*16
                                        doc.paragraphs[p].style.font.name='Times New Roman'
                                        doc.paragraphs[p].style.font.bold=True
                                        print(ttext)
        # for t in  range(0,len(doc.tables)):
        #         for r in range(0,len(doc.tables[t].rows)):
        #                 for c in range(0,len(doc.tables[t].rows[r].cells)):
        #                         stext=doc.tables[t].rows[r].cells[c].paragraphs[0].text.split()
        #                         print(stext)
        #                         rtext=''
        #                         for st in stext:
        #                                 if(st=='S.Türkmenbaşy'):
        #                                         rtext+=' Saparmyrat Türkmenbaşy '
        #                                 else:
        #                                         rtext+=' '+st
        #                         doc.tables[t].rows[r].cells[c].paragraphs[0].text=rtext
        #                         doc.tables[t].rows[r].cells[c].paragraphs[0].style.font.size=12700*16
        #                         # doc.tables[t].rows[r].cells[c].paragraphs[0]style.font.bold=True
        #                         doc.tables[t].rows[r].cells[c].paragraphs[0].style.font.name='Times New Roman'
        # for r in range(0,len(doc.tables[dtables-1].rows)):
        #         if(re.search(r'etrabynyň .* geňeşligi',doc.tables[dtables-1].rows[r].cells[3].text)):
        #                 ctext=doc.tables[dtables-1].rows[r].cells[3].text.split()
        #                 if(ctext[0]!='Garagalpagystan'):
        #                         doc.tables[dtables-1].rows[r].cells[3].text=ctext[0]+' '+ctext[1]+' '+ctext[2]+' '+'etraby'
        #                         doc.tables[dtables-1].rows[r].cells[3].paragraphs[0].style.font.size=12700*16
        #                         # doc.tables[dtables-1].rows[r].cells[3].paragraphs[0].style.font.bold=True
        #                         doc.tables[dtables-1].rows[r].cells[3].paragraphs[0].style.font.name='Times New Roman'
        #                         print(doc.tables[dtables-1].rows[r].cells[3].text)
        #         date=re.search(r'\d{4}',doc.tables[dtables-1].rows[r].cells[5].text)
        #         if(date):
        #                 ctext=doc.tables[dtables-1].rows[r].cells[5].text.split()
        #                 for ct  in range(0,len(ctext)):
        #                         if(ctext[ct].endswith('ýyl')):
        #                                 doc.tables[dtables-1].rows[r].cells[5].text=doc.tables[dtables-1].rows[r].cells[5].text[:-3]+'.'
        #                                 print(doc.tables[dtables-1].rows[r].cells[5].text[:-3])
                        # if(ctext[0]!='Garagalpagystan'):
                        #         doc.tables[dtables-1].rows[r].cells[5].text=
                        #         doc.tables[dtables-1].rows[r].cells[5].paragraphs[0].style.font.size=12700*16
                        #         # doc.tables[dtables-1].rows[r].cells[5].paragraphs[0].style.font.bold=True
                        #         doc.tables[dtables-1].rows[r].cells[5].paragraphs[0].style.font.name='Times New Roman'
                                        
                        # print(doc.tables[dtables-1].rows[r].cells[5].text)
        doc.save('prcssd/'+f)       