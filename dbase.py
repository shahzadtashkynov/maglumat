import os
import docx
import docxpy
import datetime

import mysql.connector
mydb = mysql.connector.connect(
host="localhost",
user="root",
passwd="321478569",
database="kadrlar")
cursor= mydb.cursor()

files=[]
for f in os.listdir('docs'):
    if(f.endswith('.docx')):
        files.append(f)
for f in files:
    doc=docx.Document('docs/'+f)
    table = doc.tables[0]
    data=[]
    keys='abcdefghij'
    for row in table.rows:
        text=(cell.text for cell in row.cells)
        data.append(dict(zip(keys,text)))
    for d in data:
        if(len(d)>2):
            if(d['a']==d['b']):
                d.pop('a')
            if(d['b']==d['c']):
                d.pop('c')
    # for k in data:
    #     print(k)
    # if(len(data[0]['b'].split(' '))>1):
    #     full=data[0]['b'].split(' ')
    # else:
    #     full=data[0]['b'].split(' ')
    # wsname=full[0]
    # wname=full[1]   
    # if(len(full)<3):
    #     wfname=""
    # else:
    #     wfname=full[2]
    # print (full)
    # wjob=data[1]['b'].split('mekde')
    # wbirthday=data[2]['b'].split('ý.')
    # birthday=wbirthday[0]
    # wbirthplace=wbirthday[1]
    # worker={
    #     "surname":wsname,
    #     "name":wname,
    #     "fname":wfname,
    #     "job":data[1]['b'],
    #     "working_place":wjob[0]+"mekdebi",
    #     "birthday": birthday,
    #     "birth_place":wbirthplace,
    #      "nation": data[3]['a'].split(':')[1]    ,
    #     "bilimi": data[3]['b'].split(':')[1]         ,
    #     "univer":      data[4]['b']     ,
    #     "profession":   data[5]['b']     ,
    #     "langs":  data[7]['b']               ,
    #     "derejesi":data[6]['b']            ,
    #     "dashary": data[9]['b']            ,
    #     "grants":  data[8]['b']           ,
    #     "mejlis":  data[10]['b']              ,
    #     "partiya": data[11]['b']              ,
    #     "phone":   "+99362036020"                  
    #     }
    # doc.save('prcssd/doc_'+f)
    # sql="INSERT INTO workers (imgpath,name,surname,fname, job, working_place, birthday, birth_place, nation,bilimi,univer,profession,langs,derejesi,dashary,grants,mejlis,partiya,phone) \
    # VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
    # val=("img/apelsin.jpg",worker['name'],worker['surname'],worker['fname'],worker['job'],worker['working_place'],worker['birthday'],
    # worker['birth_place'],worker['nation'],worker['bilimi'],worker['univer'],worker['profession'],worker['langs'],
    # worker['derejesi'],worker['dashary'],worker['grants'],worker['mejlis'],worker['partiya'],worker['phone'])
    # user =cursor.execute(sql,val)
    # mydb.commit()
    # print(cursor.rowcount," record inserted")
    # # reading  workers working places 
    # wid = cursor.lastrowid
    # table = doc.tables[1]
    # wplaces=[]
    # keys='abcdefghij'
    # for i,row in enumerate(table.rows):
    #     text=(cell.text for cell in row.cells)
    #     rdata = dict(zip(keys,text))
    #     wplaces.append(rdata)
    # for wp in wplaces:
    #     if ('a' or 'b' not in wp.keys()):
    #         break
    #     else:
    #         if('c' not in wp.keys()):
    #             print (wp['b'])
    #             if('-' in wp['a']):
    #                 wdate=wp['a'].split('-')
    #             if('ý.' in wp['a']):
    #                 wdate=wp['a'].split('ý.')
    #             else:
    #                 wdate=wp['a'].split('ý')
    #             wbdate=wdate[0]
    #             wedate=wdate[1]
    #             wnirede=wp['b']
    #         else:
    #             print (wp['c'])
    #             wbdate=wp['a']
    #             wedate=wp['b']
    #             wnirede=wp['c']
    #         sql="INSERT INTO worked_places (worker_id,bdate,edate,nirede)\
    #         VALUES (%s,%s,%s,%s)"
    #         val=(wid,wbdate,wedate,wnirede)
    #         cursor.execute(sql,val)
    #         mydb.commit()
    #         print(cursor.rowcount," record inserted")
